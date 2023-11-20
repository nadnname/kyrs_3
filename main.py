from PyQt5 import QtCore, QtGui, QtWidgets
from PyQt5.QtWidgets import QMessageBox,QDialog
from mysql.connector import *
from datetime import date
from random import randint
from docx import Document
import sys

# создание переменной для передачи инфы внутри окон
inf_g =[]
#класс первого окна
class Ui_MainWindow(object):
    # авторизация/регистрация
    def avtoriz_bd(self):

        #переменная для pyqt
        dialog = QDialog()

        #запись ответов
        passw=self.line2.text()
        fio = self.line1.text()
        #обработка ответов
        cur.execute('select * from klient')
        db = cur.fetchall()
        passw_check = []
        fio_check = []
        for i in range(len(db)):
            passw_check.append(db[i][1])
            fio_check.append(db[i][2])

        # Предотвращение ошибок
        if (fio == '' or fio == ' ' or fio ==None) or (passw == '' or passw == ' ' or passw == None):
            self.result = QMessageBox.critical(dialog, "Ошибка", "Присутствуют пустые поля")
        elif fio in fio_check:
                z = fio_check.index(fio)
                if passw == passw_check[z]:
                    self.result = QMessageBox.information(dialog, "Авторизация", "Авторизация прошла успешно")
                    inf_g.append(z)
                    inf_g.append(fio)
                    self.open_dog_wind()
                else:
                    self.result = QMessageBox.critical(dialog, "Ошибка", "Пароль не верный")
        else:
            self.result = QMessageBox.information(dialog, "Регистрация", "Новый пользователь зарегистрирован")
            self.zapl_bd()
            self.open_dog_wind()
    #регистрация пользователя и запись в бд
    def zapl_bd(self):
        # запись ответов в окне
        passw = self.line2.text()
        fio = self.line1.text()
        #подключение к бд и создание запроса на запись
        cur.execute('select * from klient')
        db = cur.fetchall()
        db = str(len(db) + 1)
        inf = (db, passw, fio)
        inf_g.append(db)
        inf_g.append(fio)
        zapr = ('insert into klient (id,passw,fio) values (%s,%s,%s)')
        cur.execute(zapr, inf)
        con.commit()

    #Открытие дочернего окна
    def open_dog_wind(self):
        MainWindow.close()
        # открытие второго окна
        self.window = QtWidgets.QMainWindow()
        self.ui = Ui_Form()
        self.ui.setupUi(self.window)
        self.window.show()

    #настройка виджетов и стилей главного окна
    def setupUi(self, MainWindow):
        MainWindow.setObjectName("MainWindow")
        MainWindow.resize(350, 300)
        MainWindow.setStyleSheet("background-color: rgb(217, 206, 189);\n"
"")
        MainWindow.setDocumentMode(False)
        self.centralwidget = QtWidgets.QWidget(MainWindow)
        self.centralwidget.setObjectName("centralwidget")
        self.lab1 = QtWidgets.QLabel(self.centralwidget)
        self.lab1.setGeometry(QtCore.QRect(10, 10, 271, 41))
        font = QtGui.QFont()
        font.setFamily("Arial")
        font.setPointSize(14)
        font.setBold(True)
        font.setItalic(False)
        font.setWeight(75)
        self.lab1.setFont(font)
        self.lab1.setObjectName("lab1")
        self.lab2 = QtWidgets.QLabel(self.centralwidget)
        self.lab2.setGeometry(QtCore.QRect(10, 60, 261, 21))
        font = QtGui.QFont()
        font.setFamily("Arial")
        font.setPointSize(11)
        self.lab2.setFont(font)
        self.lab2.setObjectName("lab2")
        self.but1 = QtWidgets.QPushButton(self.centralwidget)
        self.but1.clicked.connect(self.avtoriz_bd)
        self.but1.setGeometry(QtCore.QRect(10, 220, 330, 70))
        font = QtGui.QFont()
        font.setFamily("Arial")
        font.setPointSize(11)
        self.but1.setFont(font)
        self.but1.setCursor(QtGui.QCursor(QtCore.Qt.PointingHandCursor))
        self.but1.setStyleSheet("background-color: rgb(0, 170, 127);\n"
"color: rgb(255, 255, 255);\n"
"border-color: rgb(255, 255, 255);")
        self.but1.setObjectName("but1")
        self.line1 = QtWidgets.QLineEdit(self.centralwidget)
        self.line1.setGeometry(QtCore.QRect(10, 90, 171, 40))
        font = QtGui.QFont()
        font.setFamily("Arial")
        font.setPointSize(11)
        self.line1.setFont(font)
        self.line1.setStyleSheet("background-color: rgb(221, 221, 221);")
        self.line1.setObjectName("line1")
        self.line2 = QtWidgets.QLineEdit(self.centralwidget)
        self.line2.setGeometry(QtCore.QRect(10, 170, 171, 40))
        font = QtGui.QFont()
        font.setFamily("Arial")
        font.setPointSize(11)
        self.line2.setFont(font)
        self.line2.setStyleSheet("background-color: rgb(221, 221, 221);")
        self.line2.setText("")
        self.line2.setObjectName("line2")
        self.lab2_2 = QtWidgets.QLabel(self.centralwidget)
        self.lab2_2.setGeometry(QtCore.QRect(10, 140, 261, 21))
        font = QtGui.QFont()
        font.setFamily("Arial")
        font.setPointSize(11)
        self.lab2_2.setFont(font)
        self.lab2_2.setObjectName("lab2_2")
        MainWindow.setCentralWidget(self.centralwidget)

        self.retranslateUi(MainWindow)
        QtCore.QMetaObject.connectSlotsByName(MainWindow)

    #начальные данные в виджетах
    def retranslateUi(self, MainWindow):
        _translate = QtCore.QCoreApplication.translate
        MainWindow.setWindowTitle(_translate("MainWindow", "Клиент"))
        self.lab1.setText(_translate("MainWindow", "Добро пожаловать!"))
        self.lab2.setText(_translate("MainWindow", "Введите ваше ФИО:"))
        self.but1.setText(_translate("MainWindow", "Авторизация"))
        self.line1.setText(_translate("MainWindow", "Иванов И.И."))
        self.lab2_2.setText(_translate("MainWindow", "Введите ваш пароль:"))



#класс дочернего окна
class Ui_Form(object):
    #функция для всплывающего окна
    def messgbx_pyst(self):
        msg = QMessageBox()
        msg.setWindowTitle('Ошибка')
        msg.setText('Присутствуют пустые строки')
        msg.setIcon(QMessageBox.Critical)
        x = msg.exec_()

    #функция для всплывающего окна
    def messgb_docx(self):
        msg = QMessageBox()
        msg.setWindowTitle('Информация')
        msg.setText('Договор заключен, откройте его')
        msg.setIcon(QMessageBox.Information)
        x = msg.exec_()
        Form.close()

    #функция для всплывающего окна
    def messgbx_incor(self):
        msg = QMessageBox()
        msg.setWindowTitle('Ошибка')
        msg.setText('Некоректно введены данные')
        msg.setIcon(QMessageBox.Critical)
        x = msg.exec_()

    #функция для всплывающего окна
    def messgbx_exis(self):
        msg = QMessageBox()
        msg.setWindowTitle('Информация')
        msg.setText('Данные о квартире имеются в БД')
        msg.setIcon(QMessageBox.Information)
        x = msg.exec_()

    #функция для обработки данных введенных в дочернем окне
    def check_dog(self):
        #получение данных
        tip = self.com2.currentText()
        adress = self.line2.text()
        sqr = str(self.line3.text())
        stoimost = str(self.line4.text())
        #обработка данных
        cur.execute('select * from nedviz')
        db = cur.fetchall()
        adress_check = []
        tip_check = []
        sqr_check = []
        for i in db:
            sqr_check.append(str(i[1]))
            tip_check.append(i[2])
            adress_check.append(i[4])
        #добавление исключений и предотвращение ошибок
        if ((adress == '' or adress == ' ' or adress == None) or (sqr == '' or sqr == ' ' or sqr == None or int(sqr) <= 20 or int(sqr) >= 150) or (self.radio1.isChecked() == False and self.radio2.isChecked() == False)):
                self.messgbx_pyst()
        else:
            if adress in adress_check:
                p = adress_check.index(adress)
                if sqr == sqr_check[p] and tip == tip_check[p]:
                        self.messgbx_exis()
                else:
                        self.messgbx_incor()
            else:
                #если всё правильно, создаем запись в БД и заключаем договор
                if self.radio1.isChecked() == True and self.radio2.isChecked() == False and stoimost != 'Не оценено':
                        inf = (len(db)+1, sqr, tip, stoimost, adress)
                        zapr = ('insert into nedviz (id,plosh,kol_vo_komnat,stoimost,adress) values (%s,%s,%s,%s,%s)')
                        cur.execute(zapr, inf)
                        con.commit()
                        self.zakl_dog()
                elif self.radio1.isChecked() == False and self.radio2.isChecked() == True and stoimost=='На оцененке':
                        inf = (len(db) + 1, sqr, tip, stoimost, adress)
                        zapr = ('insert into nedviz (id,plosh,kol_vo_komnat,stoimost,adress) values (%s,%s,%s,%s,%s)')
                        cur.execute(zapr, inf)
                        con.commit()
                        self.zakl_dog()

                else:
                    self.messgbx_incor()
    #функция заключения договора
    def zakl_dog(self):
        #обработка ответов
        tip = self.com2.currentText()
        adress = self.line2.text()
        cur.execute('select max(nomer) from dogovor')
        db = cur.fetchone()
        cur.execute(f"select id from nedviz where adress = '{adress}'")
        nedviz_id = cur.fetchone()
        cur.execute('select max(id) from sotrydnik')
        kol_sotr = cur.fetchone()
        if tip == 'Однокомнатная':
            cymma = 10000
        elif tip == 'Двухкомнатная':
            cymma = 15000
        elif tip == 'Трехкомнатная':
            cymma = 20000
        else:
            cymma = 30000
        if self.radio1.isChecked() == True:
            tip_ysl = 'Продажа'
        else:
            tip_ysl = 'Оценка'
        #создание договора
        inf = []
        inf.append(db[0] + 1)
        inf.append(tip_ysl)
        inf.append(date.today())
        inf.append(cymma)
        inf.append(inf_g[0])
        inf.append(nedviz_id[0])
        inf.append(randint(1, kol_sotr[0]))
        zapr=('insert into dogovor(nomer, tip_yslygi, data_zakl, cymma, klient_id, nedviz_id, sotrydnik_id) values (%s,%s,%s,%s,%s,%s,%s)')
        cur.execute(zapr, inf)
        con.commit()
        #формирование выходного docx документа
        inf = []
        inf.append(str(cymma))
        if tip_ysl == 'Продажа':
            inf.append('Продажи, ')
        else:
            inf.append('Оценки, ')
        inf.append(inf_g[1])
        cur.execute(f"select * from nedviz where id = '{nedviz_id[0]}'")
        a = cur.fetchone()
        inf.append(a[4])
        inf.append(a[2])
        inf.append(str(a[1]))
        inf.append(a[3])
        inf.append(date.today())
        document = Document()
        document.add_heading('Договор', 0)
        document.add_heading('На сумму: ' + inf[0] + ' рублей.', level=1)
        document.add_heading('Об окзании услуг - ' + inf[1] + 'клиенту ' + inf[2] + ' по квартире:', level=1)
        document.add_paragraph('Адресс квартиры: ' + inf[3], style='List Bullet')
        document.add_paragraph('Тип квартиры: ' + inf[4], style='List Bullet')
        document.add_paragraph('Площадь: ' + inf[5], style='List Bullet')
        if inf[1] == 'Продажи, ':
            document.add_paragraph('Стоимость: ' + inf[6], style='List Bullet')
        document.add_heading('Дата: ' + str(inf[7]), level=1)
        document.add_paragraph('Подпись: ')
        document.add_paragraph()
        document.add_page_break()

        document.save('dogovor.docx')
        self.messgb_docx()
    #Закрывающая функция
    def close_from(self):
        dialog = QDialog()
        self.result = QMessageBox.question(dialog, "Уточнение", "Вы уверены, что хотите продолжить?")
        if self.result == QMessageBox.Yes:
            Form.close()
    #Создание дочернего окна и его виджетов
    def setupUi(self, Form):
        Form.setObjectName("Form")
        Form.resize(500, 700)
        Form.setStyleSheet("background-color: rgb(217, 206, 189);\n"
"font: 11pt \"Arial\";\n"
"")
        self.com2 = QtWidgets.QComboBox(Form)
        self.com2.setGeometry(QtCore.QRect(170, 50, 220, 30))
        self.com2.setCursor(QtGui.QCursor(QtCore.Qt.PointingHandCursor))
        self.com2.setAccessibleName("")
        self.com2.setAccessibleDescription("")
        self.com2.setStyleSheet("background-color: rgb(221, 221, 221);")
        self.com2.setObjectName("com2")
        self.com2.addItem("")
        self.com2.addItem("")
        self.com2.addItem("")
        self.com2.addItem("")
        self.lab3 = QtWidgets.QLabel(Form)
        self.lab3.setGeometry(QtCore.QRect(20, 90, 231, 21))
        self.lab3.setObjectName("lab3")
        self.lab2 = QtWidgets.QLabel(Form)
        self.lab2.setGeometry(QtCore.QRect(20, 50, 150, 30))
        self.lab2.setObjectName("lab2")
        self.groupBox = QtWidgets.QGroupBox(Form)
        self.groupBox.setGeometry(QtCore.QRect(20, 250, 461, 111))
        self.groupBox.setCursor(QtGui.QCursor(QtCore.Qt.PointingHandCursor))
        self.groupBox.setFocusPolicy(QtCore.Qt.NoFocus)
        self.groupBox.setContextMenuPolicy(QtCore.Qt.NoContextMenu)
        self.groupBox.setAcceptDrops(False)
        self.groupBox.setObjectName("groupBox")
        self.radio1 = QtWidgets.QRadioButton(self.groupBox)
        self.radio1.setGeometry(QtCore.QRect(30, 50, 131, 31))
        self.radio1.setFocusPolicy(QtCore.Qt.NoFocus)
        self.radio1.setObjectName("radio1")
        self.radio2 = QtWidgets.QRadioButton(self.groupBox)
        self.radio2.setGeometry(QtCore.QRect(270, 50, 131, 31))
        self.radio2.setObjectName("radio2")
        self.lab1 = QtWidgets.QLabel(Form)
        self.lab1.setGeometry(QtCore.QRect(20, 10, 380, 30))
        self.lab1.setObjectName("lab1")
        self.lab4 = QtWidgets.QLabel(Form)
        self.lab4.setGeometry(QtCore.QRect(20, 160, 401, 31))
        self.lab4.setObjectName("lab4")
        self.but1 = QtWidgets.QPushButton(Form)
        self.but1.clicked.connect(self.check_dog)
        self.but1.setGeometry(QtCore.QRect(20, 470, 460, 100))
        self.but1.setCursor(QtGui.QCursor(QtCore.Qt.PointingHandCursor))
        self.but1.setStyleSheet("color: rgb(255, 255, 255);\n"
"background-color: rgb(0, 170, 127);\n"
"")
        self.but1.setObjectName("but1")
        self.but2 = QtWidgets.QPushButton(Form, clicked=lambda: self.close_from())
        self.but2.setGeometry(QtCore.QRect(20, 590, 460, 100))
        self.but2.setCursor(QtGui.QCursor(QtCore.Qt.PointingHandCursor))
        self.but2.setStyleSheet("color: rgb(255, 255, 255);\n"
"background-color: rgb(170, 0, 0);")
        self.but2.setObjectName("but2")
        self.line2 = QtWidgets.QLineEdit(Form)
        self.line2.setGeometry(QtCore.QRect(20, 120, 400, 40))
        self.line2.setStyleSheet("background-color: rgb(221, 221, 221);")
        self.line2.setObjectName("line2")
        self.line3 = QtWidgets.QLineEdit(Form)
        self.line3.setGeometry(QtCore.QRect(20, 200, 250, 40))
        self.line3.setStyleSheet("background-color: rgb(221, 221, 221);")
        self.line3.setObjectName("line3")
        self.lab5 = QtWidgets.QLabel(Form)
        self.lab5.setGeometry(QtCore.QRect(20, 380, 401, 31))
        self.lab5.setObjectName("lab5")
        self.line4 = QtWidgets.QLineEdit(Form)
        self.line4.setGeometry(QtCore.QRect(20, 420, 250, 40))
        self.line4.setStyleSheet("background-color: rgb(221, 221, 221);")
        self.line4.setObjectName("line4")

        self.retranslateUi(Form)
        QtCore.QMetaObject.connectSlotsByName(Form)
    #создание начальных значений
    def retranslateUi(self, Form):
        _translate = QtCore.QCoreApplication.translate
        Form.setWindowTitle(_translate("Form", "Создание Договора"))
        self.com2.setItemText(0, _translate("Form", "Однокомнатная"))
        self.com2.setItemText(1, _translate("Form", "Двухкомнатная"))
        self.com2.setItemText(2, _translate("Form", "Трехкомнатная"))
        self.com2.setItemText(3, _translate("Form", "Более трех"))
        self.lab3.setText(_translate("Form", "Адрес квартиры"))
        self.lab2.setText(_translate("Form", "Тип квартиры:"))
        self.groupBox.setTitle(_translate("Form", "Тип услуги:"))
        self.radio1.setText(_translate("Form", "Продажа"))
        self.radio2.setText(_translate("Form", "Оценка"))
        self.lab1.setText(_translate("Form", "Заполните информацию"))
        self.lab4.setText(_translate("Form", "Площадь квартиры (в км2):"))
        self.but1.setText(_translate("Form", "Заключить договор"))
        self.but2.setText(_translate("Form", "Выйти"))
        self.line2.setText(_translate("Form", "ул.Пушкина, д.41, кв.111"))
        self.lab5.setText(_translate("Form", "Стоимость квартиры(при продаже):"))
        self.line4.setText(_translate("Form", "На оцененке"))

#запуск программы, с учетом подключения к БД
if __name__ == "__main__":
    try:
        con = connect(host='localhost', password='12345', username='root', db='kyrsach')
        cur = con.cursor()
        app = QtWidgets.QApplication(sys.argv)
        MainWindow = QtWidgets.QMainWindow()
        ui = Ui_MainWindow()
        ui.setupUi(MainWindow)
        MainWindow.show()
        sys.exit(app.exec_())
        app = QtWidgets.QApplication(sys.argv)
        Form = QtWidgets.QWidget()
        ui = Ui_Form()
        ui.setupUi(Form)
    except:
        print('Проверьте подключение к БД')
